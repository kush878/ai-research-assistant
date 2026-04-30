import requests
import streamlit as st
from datetime import datetime
from rag_pipeline import process_pdf, get_relevant_chunks
from agent import plan_question, solve_subquestions, generate_final_answer
from llm import get_answer
from embeddings import get_embeddings
from vector_store import create_faiss_index
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from streamlit_oauth import OAuth2Component
import os
import sqlite3
import hashlib
import json
from database import init_db, save_chat, load_chats, delete_chat, rename_chat

def sync_current_chat():
    save_chat(
        st.session_state.user_email,
        st.session_state.current_chat,
        json.dumps(st.session_state.messages)
    )
init_db()
# ---------------- PAGE CONFIG ----------------
st.set_page_config(
    page_title="AI Research Assistant",
    page_icon="🤖",
    layout="wide"
)
import os

GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID")
GOOGLE_CLIENT_SECRET = os.getenv("GOOGLE_CLIENT_SECRET")

AUTHORIZE_URL = "https://accounts.google.com/o/oauth2/v2/auth"
TOKEN_URL = "https://oauth2.googleapis.com/token"
REFRESH_URL = TOKEN_URL
REVOKE_URL = "https://oauth2.googleapis.com/revoke"

# ---------------- LOGIN DATABASE ----------------
def create_users_table():
    init_db()
    conn = sqlite3.connect("users.db")
    c = conn.cursor()

    c.execute("""
        CREATE TABLE IF NOT EXISTS users(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT,
            email TEXT UNIQUE,
            password TEXT
        )
    """)

    conn.commit()
    conn.close()


def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()


def register_user(username, email, password):
    conn = sqlite3.connect("users.db")
    c = conn.cursor()

    try:
        c.execute(
            "INSERT INTO users(username,email,password) VALUES(?,?,?)",
            (username, email, hash_password(password))
        )
        conn.commit()
        conn.close()
        return True
    except:
        conn.close()
        return False


def login_user(email, password):
    conn = sqlite3.connect("users.db")
    c = conn.cursor()

    c.execute(
        "SELECT * FROM users WHERE email=? AND password=?",
        (email, hash_password(password))
    )

    data = c.fetchone()
    conn.close()
    return data


create_users_table()

# ---------------- CSS ----------------
st.markdown("""
<style>
.stApp{
background:linear-gradient(135deg,#050816,#0b1120,#111827);
color:white;
}
#MainMenu,footer,header{visibility:hidden;}

.main-title{
font-size:46px;
font-weight:800;
color:white;
text-shadow:0 0 14px #7c3aed;
margin-bottom:5px;
}

.sub-title{
color:#9ca3af;
font-size:17px;
margin-bottom:22px;
}

.stButton>button{
width:100%;
background:linear-gradient(90deg,#7c3aed,#2563eb);
color:white;
border:none;
border-radius:12px;
padding:9px;
font-weight:700;
}

.stButton>button:hover{
box-shadow:0 0 15px rgba(124,58,237,0.55);
}

[data-testid="stChatMessage"]{
background:rgba(255,255,255,0.04);
padding:10px;
border-radius:14px;
margin-bottom:8px;
}

.side-box{
background:rgba(255,255,255,0.04);
padding:12px;
border-radius:12px;
margin-bottom:8px;
border:1px solid rgba(255,255,255,0.05);
}

.small{
font-size:13px;
color:#d1d5db;
}
</style>
""", unsafe_allow_html=True)

# ---------------- SESSION ----------------
defaults = {
    "messages": [],
    "index": None,
    "chunks": None,
    "tool": "",
    "relevant_chunks": [],
    "show_sidebar": True,
    "chat_sessions": {},
    "current_chat": "New Chat",
    "chat_count": 1,
    "menu_open": "",
    "rename_mode": "",
    "summary_text": "",
    "logged_in": False,
    "guest_mode": False,
    "user_name": "",
    "user_email": "",
    "user_photo": ""
}

for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

if len(st.session_state.chat_sessions) == 0:
    st.session_state.chat_sessions["New Chat"] = []

# ---------------- LOGIN PAGE ----------------
if not st.session_state.logged_in:

    st.markdown("""
    <style>
    .login-box {
        max-width: 650px;
        margin: auto;
        margin-top: 40px;
        padding: 35px;
        border-radius: 24px;
        background: rgba(255,255,255,0.05);
        backdrop-filter: blur(14px);
        box-shadow: 0 0 30px rgba(124,58,237,0.25);
        border: 1px solid rgba(255,255,255,0.08);
    }

    .title-main {
        text-align:center;
        font-size:42px;
        font-weight:800;
        color:white;
        margin-bottom:8px;
    }

    .sub-main {
        text-align:center;
        color:#9ca3af;
        font-size:17px;
        margin-bottom:28px;
    }

    .feature-line {
        text-align:center;
        color:#d1d5db;
        font-size:15px;
        margin-bottom:20px;
    }
    </style>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="login-box">

    <div class="title-main">🤖 AI Research Assistant</div>

    <div class="sub-main">
    Smart PDF Chat • Summarize • Export • Multi Chat
    </div>

    <div class="feature-line">
    Upload pdf and ask anything instantly
    </div>

    </div>
    """, unsafe_allow_html=True)

    tab1, tab2 = st.tabs(["🔑 Login", "📝 Register"])

    with tab1:

        email = st.text_input("📧 Email")
        password = st.text_input("🔒 Password", type="password")

        if st.button("🚀 Login", use_container_width=True):
            user = login_user(email, password)

            if user:
                st.session_state.logged_in = True
                st.session_state.user_name = user[1]
                st.session_state.user_email = user[2]

                old = load_chats(user[2])
                st.session_state.chat_sessions = {}

                for name, msgs in old:
                    try:
                        st.session_state.chat_sessions[name] = json.loads(msgs)
                    except:
                        st.session_state.chat_sessions[name] = []

                if len(old) > 0:
                    st.session_state.current_chat = old[0][0]
                    st.session_state.messages = st.session_state.chat_sessions[old[0][0]]
                else:
                    st.session_state.chat_sessions["New Chat"] = []
                    st.session_state.current_chat = "New Chat"
                    st.session_state.messages = []

                st.rerun()

            else:
                st.error("Invalid Email or Password")

        st.markdown("###")

        oauth2 = OAuth2Component(
            GOOGLE_CLIENT_ID,
            GOOGLE_CLIENT_SECRET,
            AUTHORIZE_URL,
            TOKEN_URL,
            REFRESH_URL,
            REVOKE_URL
        )

        result = oauth2.authorize_button(
            name="🔵 Continue with Google",
            icon="https://www.google.com/favicon.ico",
            redirect_uri="http://localhost:8501",
            scope="openid email profile",
            key="google_login"
        )

        if result:
            access_token = result["token"]["access_token"]

            user_info = requests.get(
                "https://www.googleapis.com/oauth2/v1/userinfo",
                params={"alt": "json"},
                headers={"Authorization": f"Bearer {access_token}"}
            ).json()

            real_name = user_info.get("name", "Google User")
            real_email = user_info.get("email", "unknown@gmail.com")
            real_photo = user_info.get("picture", "")

            st.session_state.logged_in = True
            st.session_state.user_name = real_name
            st.session_state.user_email = real_email
            st.session_state.user_photo = real_photo

            old = load_chats(real_email)
            st.session_state.chat_sessions = {}

            for name, msgs in old:
                try:
                    st.session_state.chat_sessions[name] = json.loads(msgs)
                except:
                    st.session_state.chat_sessions[name] = []

            if len(old) > 0:
                st.session_state.current_chat = old[0][0]
                st.session_state.messages = st.session_state.chat_sessions[old[0][0]]
            else:
                st.session_state.chat_sessions["New Chat"] = []
                st.session_state.current_chat = "New Chat"
                st.session_state.messages = []

            st.rerun()

        if st.button("👤 Continue as Guest", use_container_width=True):
            st.session_state.logged_in = True
            st.session_state.user_name = "Guest User"
            st.session_state.guest_mode = True
            st.rerun()

    with tab2:

        username = st.text_input("👤 Username")
        reg_email = st.text_input("📧 Register Email")
        reg_pass = st.text_input("🔒 Create Password", type="password")

        if st.button("📝 Create Account", use_container_width=True):
            ok = register_user(username, reg_email, reg_pass)

            if ok:
                st.success("Registration Successful")
            else:
                st.error("Email already exists.")

    st.stop()
# # ---------------- LOGIN PAGE ----------------
# if not st.session_state.logged_in:

#     st.markdown("""
#     <h1 style='text-align:center;color:white;'>🔐 AI Research Assistant</h1>
#     <h4 style='text-align:center;color:#9ca3af;'>Login to Continue</h4>
#     """, unsafe_allow_html=True)

#     tab1, tab2 = st.tabs(["🔑 Login", "📝 Register"])

#     with tab1:

#         email = st.text_input("Email")
#         password = st.text_input("Password", type="password")

#         if st.button("Login"):
#             user = login_user(email, password)

#             if user:
#                 st.session_state.logged_in = True
#                 st.session_state.user_name = user[1]
#                 st.session_state.user_email = user[2]
                
#                 old = load_chats(user[2])
#                 st.session_state.chat_sessions = {}
                
#                 for name, msgs in old:
#                     try:
#                         st.session_state.chat_sessions[name] = json.loads(msgs)
#                     except:
#                         st.session_state.chat_sessions[name] = []
                        
#                 if len(old) > 0:
#                     st.session_state.current_chat = old[0][0]
#                     st.session_state.messages = st.session_state.chat_sessions[old[0][0]]
#                 else:
#                     st.session_state.chat_sessions["New Chat"] = []
#                     st.session_state.current_chat = "New Chat"
#                     st.session_state.messages = []
#                     st.rerun()
#                 st.rerun()
#             else:
#                 st.error("Invalid Email or Password")

#         st.markdown("### OR")

#         oauth2 = OAuth2Component(
#             GOOGLE_CLIENT_ID,
#             GOOGLE_CLIENT_SECRET,
#             AUTHORIZE_URL,
#             TOKEN_URL,
#             REFRESH_URL,
#             REVOKE_URL
#         )

#         result = oauth2.authorize_button(
#             name="🔵 Continue with Google",
#             icon="https://www.google.com/favicon.ico",
#             redirect_uri="http://localhost:8501",
#             scope="openid email profile",
#             key="google_login"
#         )

#         if result:
#             access_token = result["token"]["access_token"]
            
#             user_info = requests.get("https://www.googleapis.com/oauth2/v1/userinfo",
#                                      params={"alt": "json"},
#                                      headers={"Authorization": f"Bearer {access_token}"}
#                                      ).json()
            
#             real_name = user_info.get("name", "Google User")
#             real_email = user_info.get("email", "unknown@gmail.com")
#             real_photo = user_info.get("picture", "")
            
#             st.session_state.logged_in = True
#             st.session_state.user_name = real_name
#             st.session_state.user_email = real_email
#             st.session_state.user_photo = real_photo
            
#             old = load_chats(real_email)
#             st.session_state.chat_sessions = {}
#             for name, msgs in old:
#                 try:
#                     st.session_state.chat_sessions[name] = json.loads(msgs)
#                 except:
#                     st.session_state.chat_sessions[name] = []
                    
#             if len(old) > 0:
#                 st.session_state.current_chat = old[0][0]
#                 st.session_state.messages = st.session_state.chat_sessions[old[0][0]]
#             else:
#                 st.session_state.chat_sessions["New Chat"] = []
#                 st.session_state.current_chat = "New Chat"
#                 st.session_state.messages = []
#             st.rerun()

#         if st.button("👤 Continue as Guest"):
#             st.session_state.logged_in = True
#             st.session_state.user_name = "Guest User"
#             st.session_state.guest_mode = True
#             st.rerun()

#     with tab2:

#         username = st.text_input("Username")
#         reg_email = st.text_input("Register Email")
#         reg_pass = st.text_input("Create Password", type="password")

#         if st.button("Register"):
#             ok = register_user(username, reg_email, reg_pass)

#             if ok:
#                 st.success("Registration Successful")
#             else:
#                 st.error("Email already exists.")

#     st.stop()
# ---------------- EXPORT FUNCTIONS ----------------
def generate_word_file(chat_name, summary, messages):
    buffer = BytesIO()
    doc = Document()

    doc.add_heading("AI Research Paper Assistant Report", 0)
    doc.add_paragraph(f"Session Name: {chat_name}")
    doc.add_paragraph(f"Generated On: {datetime.now().strftime('%d-%m-%Y %H:%M')}")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary if summary else "No summary generated.")

    doc.add_heading("Chat History", level=1)

    if messages:
        for msg in messages:
            doc.add_paragraph(f"{msg['role'].upper()}: {msg['content']}")
    else:
        doc.add_paragraph("No chat available.")

    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_pdf_file(chat_name, summary, messages):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)

    width, height = A4
    y = height - 40

    c.setFont("Helvetica-Bold", 16)
    c.drawString(40, y, "AI Research Paper Assistant Report")

    y -= 25
    c.setFont("Helvetica", 10)
    c.drawString(40, y, f"Session Name: {chat_name}")

    y -= 18
    c.drawString(40, y, f"Generated On: {datetime.now().strftime('%d-%m-%Y %H:%M')}")

    y -= 30
    c.setFont("Helvetica-Bold", 12)
    c.drawString(40, y, "Summary")

    y -= 18
    c.setFont("Helvetica", 10)

    summary_text = summary if summary else "No summary generated."

    for line in summary_text.split("\n"):
        c.drawString(40, y, line[:100])
        y -= 14

    y -= 20
    c.setFont("Helvetica-Bold", 12)
    c.drawString(40, y, "Chat History")

    y -= 18
    c.setFont("Helvetica", 10)

    if messages:
        for msg in messages:
            txt = f"{msg['role'].upper()}: {msg['content']}"
            parts = [txt[i:i+100] for i in range(0, len(txt), 100)]

            for p in parts:
                c.drawString(40, y, p)
                y -= 14

                if y < 40:
                    c.showPage()
                    y = height - 40

    c.save()
    buffer.seek(0)
    return buffer

# ---------------- TOGGLE ----------------
top1, top2 = st.columns([1,14])

with top1:
    icon = "<<" if st.session_state.show_sidebar else ">>"
    if st.button(icon):
        st.session_state.show_sidebar = not st.session_state.show_sidebar
        st.rerun()

# ---------------- LAYOUT ----------------
if st.session_state.show_sidebar:
    left, right = st.columns([1,4])
else:
    left, right = st.columns([0.01,5])

# ---------------- SIDEBAR ----------------
with left:
    if st.session_state.show_sidebar:

        st.markdown("## 🤖 AI Assistant")
        if st.session_state.user_photo:
            st.image(st.session_state.user_photo, width=70)
            
            st.success(f"👋 Welcome {st.session_state.user_name}")
            st.caption(st.session_state.user_email)
        if st.session_state.guest_mode:
            st.info("👤 Guest Mode Active")

        if st.button("🚪 Logout"):
            st.session_state.logged_in = False
            st.session_state.user_name = ""
            st.session_state.user_email = ""
            st.session_state.user_photo = ""
            st.session_state.messages = []
            st.session_state.chat_sessions = {"New Chat": []}
            st.session_state.current_chat = "New Chat"
            st.session_state.summary_text = ""
            st.session_state.relevant_chunks = []
            st.rerun()

        if st.button("➕ New Chat"):
            st.session_state.chat_count += 1
            new_name = f"New Chat {st.session_state.chat_count}"
            
            if "chat_sessions" not in st.session_state:
                st.session_state.chat_sessions = {}
                
            st.session_state.chat_sessions[new_name] = []
            st.session_state.current_chat = new_name
            st.session_state.messages = []
            st.session_state.summary_text = ""
            
            save_chat(
                st.session_state.user_email,
                new_name,
                json.dumps([])
                )
            st.rerun()

        st.markdown("### 💬 Recent Chats")

        for chat_name in reversed(list(st.session_state.chat_sessions.keys())):

            row1, row2 = st.columns([6,1])
            short_name = chat_name[:28] + "..." if len(chat_name) > 28 else chat_name

            with row1:
                if st.button(
                    f"📄 {short_name}",
                    key=f"open_{chat_name}",
                    use_container_width=True
                ):
                    st.session_state.current_chat = chat_name
                    old = load_chats(st.session_state.user_email)
                    
                    loaded_msgs = []
                    for name, msgs in old:
                        if name == chat_name:
                            try:
                                loaded_msgs = json.loads(msgs)
                            except:
                                loaded_msgs = []
                            break
                        
                    st.session_state.messages = loaded_msgs
                    # -------- AUTO LOAD SAVED PDF --------
                    #Reset previous PDF memory first
                    st.session_state.index = None
                    st.session_state.chunks = None
                    chat_folder = os.path.join(
                        "uploads",
                        st.session_state.user_email.replace("@", "_"),
                        chat_name
                        )
                    
                    if os.path.exists(chat_folder):
                        pdf_files = [
                            f for f in os.listdir(chat_folder)
                            if f.lower().endswith(".pdf")
                            ]
                        if len(pdf_files) > 0:
                            all_chunks = []
                            for pdf_name in pdf_files:
                                pdf_path = os.path.join(chat_folder, pdf_name)
                                with open(pdf_path, "rb") as f:
                                    index, chunks = process_pdf(f)
                                    all_chunks.extend(chunks)
                                    embeddings = get_embeddings(all_chunks)
                                    st.session_state.index = create_faiss_index(embeddings)
                                    st.session_state.chunks = all_chunks
                    st.session_state.chat_sessions[chat_name] = loaded_msgs
                    st.session_state.relevant_chunks = []
                    st.session_state.summary_text = ""
                    st.rerun()

            with row2:
                if st.button("⋮", key=f"menu_{chat_name}", use_container_width=True):
                    if st.session_state.menu_open == chat_name:
                        st.session_state.menu_open = ""
                    else:
                        st.session_state.menu_open = chat_name

            if st.session_state.menu_open == chat_name:

                m1, m2 = st.columns(2)

                with m1:
                    if st.button("✏️ Rename", key=f"rename_{chat_name}"):
                        st.session_state.rename_mode = chat_name

                with m2:
                    if st.button("🗑 Delete", key=f"delete_{chat_name}"):
                        delete_chat(
                            st.session_state.user_email,
                            chat_name
                            )
                        
                        del st.session_state.chat_sessions[chat_name]

                        if len(st.session_state.chat_sessions) == 0:
                            st.session_state.chat_sessions["New Chat"] = []
                            st.session_state.current_chat = "New Chat"
                            st.session_state.messages = []
                        else:
                            first_chat = list(st.session_state.chat_sessions.keys())[0]
                            st.session_state.current_chat = first_chat
                            st.session_state.messages = st.session_state.chat_sessions[first_chat]

                        st.rerun()

            if st.session_state.rename_mode == chat_name:

                new_title = st.text_input(
                    "Rename Chat",
                    value=chat_name,
                    key=f"rename_input_{chat_name}"
                )

                if st.button("Save", key=f"save_{chat_name}"):
                    rename_chat(
                        st.session_state.user_email,
                        chat_name,
                        new_title
                        )

                    if new_title.strip() != "" and new_title not in st.session_state.chat_sessions:
                        data = st.session_state.chat_sessions.pop(chat_name)
                        st.session_state.chat_sessions[new_title] = data
                        save_chat(
                            st.session_state.user_email,
                            new_title,
                            json.dumps(data)
                            )

                        if st.session_state.current_chat == chat_name:
                            st.session_state.current_chat = new_title

                    st.session_state.rename_mode = ""
                    st.rerun()

        st.markdown("---")

        # total_chunks = len(st.session_state.chunks) if st.session_state.chunks else 0

        # st.markdown(f"""
        # <div class="side-box">
        # 🚀 <b>System</b><br><br>
        # <span class="small">
        # 🧠 Agentic RAG<br>
        # ⚡ Groq API<br>
        # 📅 {datetime.now().strftime("%d-%m-%Y")}<br>
        # ⏰ {datetime.now().strftime("%H:%M")}
        # </span>
        # </div>
        # """, unsafe_allow_html=True)

        # st.markdown(f"""
        # <div class="side-box">
        # 📄 <b>Documents</b><br><br>
        # <span class="small">
        # 📚 Chunks: {total_chunks}<br>
        # 📑 Ready: {"Yes" if total_chunks > 0 else "No"}
        # </span>
        # </div>
        # """, unsafe_allow_html=True)

# ---------------- MAIN ----------------
with right:

    st.markdown('<div class="main-title">🤖 AI Research Assistant</div>', unsafe_allow_html=True)

    st.markdown(
        f'<div class="sub-title">Current Session: {st.session_state.current_chat}</div>',
        unsafe_allow_html=True
    )

    uploaded_files = st.file_uploader(
        "📂 Upload Research Papers",
        type="pdf",
        accept_multiple_files=True
    )

    if uploaded_files:

        if len(uploaded_files) == 1:
            auto_name = os.path.splitext(uploaded_files[0].name)[0]
        else:
            auto_name = "Research Bundle"

        old_chat = st.session_state.current_chat

        if old_chat.startswith("New Chat"):

            if auto_name not in st.session_state.chat_sessions:
                st.session_state.chat_sessions[auto_name] = st.session_state.chat_sessions.pop(old_chat)
                st.session_state.current_chat = auto_name
                st.session_state.messages = st.session_state.chat_sessions[auto_name]

        st.success("✅ PDFs uploaded successfully!")
        import os
        
        chat_folder = os.path.join(
            "uploads",
            st.session_state.user_email.replace("@", "_"),
            st.session_state.current_chat
            )
        
        os.makedirs(chat_folder, exist_ok=True)
        for file in uploaded_files:
            save_path = os.path.join(chat_folder, file.name)
            
            with open(save_path, "wb") as f:
                f.write(file.getbuffer())

        all_chunks = []

        with st.spinner("Processing PDFs..."):
            for file in uploaded_files:
                index, chunks = process_pdf(file)
                all_chunks.extend(chunks)

            embeddings = get_embeddings(all_chunks)
            st.session_state.index = create_faiss_index(embeddings)
            st.session_state.chunks = all_chunks
            save_chat(
                st.session_state.user_email,
                st.session_state.current_chat,
                json.dumps(st.session_state.messages)
                )

    c1, c2 = st.columns(2)

    with c1:
        if st.button("📄 Summarize Document"):
            if st.session_state.chunks is not None:
                with st.spinner("Generating Summary..."):
                    text = " ".join(st.session_state.chunks[:25])
                    st.session_state.summary_text = get_answer(
                        text,
                        "Create a detailed clear summary with key points."
                        )
            else:
                st.warning("Open the chat where PDF is uploaded or upload PDF again.")

    with c2:
        if st.button("🗑 Clear Current Chat"):
            st.session_state.messages = []
            st.session_state.chat_sessions[st.session_state.current_chat] = []
            st.session_state.summary_text = ""
            
            save_chat(
                st.session_state.user_email,
                st.session_state.current_chat,
                json.dumps([])
            )
            
            st.success("Current chat cleared!")

    if st.session_state.summary_text:
        st.subheader("📌 Summary")
        st.write(st.session_state.summary_text)

    st.markdown("### 📤 Export Reports")

    ex1, ex2 = st.columns(2)

    with ex1:
        st.download_button(
            "📝 Download Word Report",
            data=generate_word_file(
                st.session_state.current_chat,
                st.session_state.summary_text,
                st.session_state.messages
            ),
            file_name=f"{st.session_state.current_chat}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    with ex2:
        st.download_button(
            "📄 Download PDF Report",
            data=generate_pdf_file(
                st.session_state.current_chat,
                st.session_state.summary_text,
                st.session_state.messages
            ),
            file_name=f"{st.session_state.current_chat}.pdf",
            mime="application/pdf",
            use_container_width=True
        )

    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.write(msg["content"])
            
    if len(st.session_state.messages) == 0:
        st.info("This chat has no saved messages yet. Upload PDF and start asking questions.")
        
    import time
    
    if "last_call_ts" not in st.session_state:
        st.session_state.last_call_ts = 0
    user_input = st.chat_input("Ask anything about the research paper...")
    if user_input:
        
        # ⏳ Cooldown check (ADD THIS BLOCK)
        now = time.time()
        COOLDOWN = 6
        
        if now - st.session_state.last_call_ts < COOLDOWN:
            wait = int(COOLDOWN - (now - st.session_state.last_call_ts)) + 1
            st.warning(f"⏳ Please wait {wait}s before next question.")
            st.stop()
            
        # EXISTING CODE CONTINUES
        if st.session_state.index is None:
            st.warning("PDF context not loaded. Open recent uploaded chat or upload PDF again.")
        else:
            st.session_state.messages.append({"role": "user", "content": user_input})
            sync_current_chat()
            
            relevant_chunks = get_relevant_chunks(
                user_input,
                st.session_state.index,
                st.session_state.chunks
            )
            context = relevant_chunks[0][:900]
            answer = get_answer(context, user_input)
            
            #important (save time for cooldown)
            st.session_state.last_call_ts = time.time()
            st.session_state.relevant_chunks = relevant_chunks
            st.session_state.messages.append({"role":"assistant","content":answer})
            sync_current_chat()

    # if user_input:

    #     if st.session_state.index is None:
    #         st.warning("PDF context not loaded. Open recent uploaded chat or upload PDF again.")
    #     else:
    #         st.session_state.messages.append({"role": "user", "content": user_input})
    #         sync_current_chat()


    #         st.session_state.relevant_chunks = relevant_chunks

    #         with st.spinner("🧠 AI Thinking..."):

    #             sub_questions = plan_question(user_input)

    #             sub_answers = solve_subquestions(
    #                 sub_questions,
    #                 st.session_state.index,
    #                 st.session_state.chunks
    #             )

    #             answer, tool = generate_final_answer(user_input, sub_answers)
    #             st.session_state.tool = tool

    #         st.session_state.messages.append(
    #             {"role": "assistant", "content": answer}
    #         )
    #         sync_current_chat()

            st.session_state.chat_sessions[
                st.session_state.current_chat
            ] = st.session_state.messages
            
            save_chat(
                st.session_state.user_email,
                st.session_state.current_chat,
                json.dumps(st.session_state.messages)
            )

            st.rerun()

    if st.session_state.relevant_chunks:
        st.subheader("📚 Source Insights")

        for i, chunk in enumerate(st.session_state.relevant_chunks[:3]):
            st.markdown(f"""
            <div class="side-box">
            <b style="color:#8b5cf6;">📄 Source {i+1}</b><br><br>
            <span class="small">
            {chunk[:550]}...
            </span>
            </div>
            """, unsafe_allow_html=True)

# # Footer
# st.markdown("---")
# st.markdown(
#     "<center style='color:#9ca3af;'>⚡ Built by Kush Vaishnani | Final Professional Agentic RAG</center>",
#     unsafe_allow_html=True
# )
